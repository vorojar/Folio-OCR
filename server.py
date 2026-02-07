# -*- coding: utf-8 -*-
"""
Folio-OCR Server - Ollama Backend
Three-column document workbench
"""
import os
import uuid
import time
import asyncio
import base64
import logging
import shutil
import subprocess
import httpx
from pathlib import Path
from datetime import datetime
from fastapi import FastAPI, UploadFile, File, HTTPException, Query
import json
from fastapi.responses import HTMLResponse, FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import re
from html.parser import HTMLParser
from urllib.parse import quote
import fitz  # PyMuPDF for PDF processing
from PIL import Image
import io
import torch
from transformers import AutoImageProcessor, AutoModelForObjectDetection
from pydantic import BaseModel
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Setup logging
LOG_FILE = Path(__file__).parent / "server.log"
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(LOG_FILE, encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)
logging.getLogger("httpx").setLevel(logging.WARNING)
logger.info(f"=== Server starting, log file: {LOG_FILE} ===")

app = FastAPI(title="Folio-OCR Service", version="3.1.0")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directories
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# Ollama config
OLLAMA_BASE = "http://localhost:11434"
OLLAMA_MODEL = "glm-ocr"
OCR_PROMPT = "识别图片中的全部内容，输出Markdown格式。跳过页眉页脚和页码。"

# LaTeX → Unicode mapping (loaded once at import time)
_LATEX_MAP_FILE = Path(__file__).parent / "latex_unicode.json"
with open(_LATEX_MAP_FILE, "r", encoding="utf-8") as _f:
    _LATEX_DATA = json.load(_f)
_LATEX_SIMPLE: list[tuple[str, str]] = sorted(
    _LATEX_DATA["simple"].items(), key=lambda x: -len(x[0])
)  # longest match first
_LATEX_FRACTIONS: dict[str, str] = _LATEX_DATA.get("fractions", {})
_CIRCLED = {str(i): chr(0x2460 + i - 1) for i in range(1, 21)}  # ①-⑳

# In-memory document state
# doc_id -> {filename, pages: [{num, filename, ocr_text, ocr_time}], created_at}
documents: dict[str, dict] = {}

# Shared httpx client
_http_client: httpx.AsyncClient | None = None

# Layout detection model (PP-DocLayoutV3)
_LAYOUT_MODEL_NAME = "PaddlePaddle/PP-DocLayoutV3_safetensors"
_layout_processor = None
_layout_model = None
# Labels to skip (not useful for OCR text mapping)
_LAYOUT_SKIP_LABELS = {"header", "footer", "footnote", "number", "footer"}
_LAYOUT_THRESHOLD = 0.5


@app.on_event("startup")
async def startup_event():
    global _http_client, _layout_processor, _layout_model
    _http_client = httpx.AsyncClient(timeout=httpx.Timeout(300.0))

    # Load layout detection model
    t0 = time.time()
    logger.info(f"[layout] Loading {_LAYOUT_MODEL_NAME}...")
    _layout_processor = AutoImageProcessor.from_pretrained(_LAYOUT_MODEL_NAME)
    _layout_model = AutoModelForObjectDetection.from_pretrained(_LAYOUT_MODEL_NAME)
    if torch.cuda.is_available():
        _layout_model.to("cuda")
        logger.info(f"[layout] Model loaded on CUDA: {time.time() - t0:.2f}s")
    else:
        logger.info(f"[layout] Model loaded on CPU: {time.time() - t0:.2f}s")
    _layout_model.eval()

    # Clean up orphan directories from previous runs
    for child in UPLOAD_DIR.iterdir():
        if child.is_dir() and child.name not in documents:
            logger.info(f"[cleanup] Removing orphan directory: {child}")
            shutil.rmtree(child, ignore_errors=True)


@app.on_event("shutdown")
async def shutdown_event():
    global _http_client
    if _http_client:
        await _http_client.aclose()


async def check_ollama() -> dict:
    """Check Ollama status and model availability"""
    try:
        resp = await _http_client.get(f"{OLLAMA_BASE}/api/tags", timeout=5.0)
        resp.raise_for_status()
        data = resp.json()
        models = [m["name"] for m in data.get("models", [])]
        has_model = any(OLLAMA_MODEL in m for m in models)
        return {"online": True, "model_loaded": has_model, "models": models}
    except Exception:
        return {"online": False, "model_loaded": False, "models": []}


# Max image height for single OCR call (fallback when layout detection returns nothing)
MAX_IMAGE_HEIGHT = 1600
SEGMENT_OVERLAP = 80


def detect_layout(img: Image.Image) -> list[dict]:
    """Detect document layout regions using PP-DocLayoutV3.
    Returns [{label, bbox: [x1,y1,x2,y2], score}] sorted by reading order (top-to-bottom).
    """
    device = next(_layout_model.parameters()).device
    inputs = _layout_processor(images=img, return_tensors="pt")
    inputs = {k: v.to(device) for k, v in inputs.items()}

    with torch.no_grad():
        outputs = _layout_model(**inputs)

    target_sizes = torch.tensor([img.size[::-1]], device=device)  # (height, width)
    results = _layout_processor.post_process_object_detection(
        outputs, target_sizes=target_sizes, threshold=_LAYOUT_THRESHOLD
    )[0]

    regions = []
    id2label = _layout_model.config.id2label
    for score, label_id, box in zip(results["scores"], results["labels"], results["boxes"]):
        label = id2label[label_id.item()]
        if label in _LAYOUT_SKIP_LABELS:
            continue
        bbox = [round(c) for c in box.tolist()]  # [x1, y1, x2, y2] in pixels
        regions.append({"label": label, "bbox": bbox, "score": round(score.item(), 3)})

    # Sort by reading order: top-to-bottom, then left-to-right
    regions.sort(key=lambda r: (r["bbox"][1], r["bbox"][0]))

    logger.info(f"[layout] Detected {len(regions)} regions")
    return regions


async def ocr_image_with_layout(image_path: str) -> tuple[str, list[dict]]:
    """Run layout detection + per-region OCR.
    Returns (combined_text, regions_list) where each region has {idx, label, bbox, text}.
    """
    t0 = time.time()
    img = Image.open(image_path).convert("RGB")

    # Step 1: Layout detection
    t1 = time.time()
    raw_regions = detect_layout(img)
    logger.info(f"[OCR] Layout detection: {time.time() - t1:.2f}s, {len(raw_regions)} regions")

    # Fallback: if no regions detected, OCR the whole image
    if not raw_regions:
        logger.info("[OCR] No layout regions, fallback to whole-image OCR")
        text = await _ocr_whole_image(img)
        img.close()
        elapsed = time.time() - t0
        logger.info(f"[OCR] TOTAL (fallback): {elapsed:.2f}s")
        return text, []

    # Step 2: Crop and OCR each region
    regions = []
    for i, region in enumerate(raw_regions):
        bbox = region["bbox"]
        cropped = img.crop(bbox)
        seg_b64 = _image_to_b64(cropped)
        text = await _ocr_single(seg_b64)
        text = _postprocess(text)

        regions.append({
            "idx": i,
            "label": region["label"],
            "bbox": bbox,
            "text": text or "",
        })
        logger.info(f"[OCR] Region {i+1}/{len(raw_regions)} ({region['label']}): {len(text)} chars")

    img.close()

    # Combine all text for backward compatibility
    combined = "\n\n".join(r["text"] for r in regions if r["text"])
    logger.info(f"[OCR] TOTAL: {time.time() - t0:.2f}s ({len(regions)} regions)")
    return combined, regions


async def _ocr_whole_image(img: Image.Image) -> str:
    """Fallback: OCR whole image, with splitting for tall images."""
    w, h = img.size
    if h > MAX_IMAGE_HEIGHT:
        segments = _split_image(img)
    else:
        segments = [img]

    all_text = []
    for seg in segments:
        seg_b64 = _image_to_b64(seg)
        text = await _ocr_single(seg_b64)
        text = _postprocess(text)
        if text:
            all_text.append(text)
    return "\n\n".join(all_text)


def _split_image(img: Image.Image) -> list[Image.Image]:
    """Split a tall image into overlapping segments."""
    w, h = img.size
    step = MAX_IMAGE_HEIGHT - SEGMENT_OVERLAP
    segments = []
    y = 0
    while y < h:
        bottom = min(y + MAX_IMAGE_HEIGHT, h)
        seg = img.crop((0, y, w, bottom))
        segments.append(seg)
        y += step
        if bottom == h:
            break
    return segments


def _image_to_b64(img: Image.Image) -> str:
    """Convert PIL Image to base64 PNG string."""
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("utf-8")


async def _ocr_single(image_b64: str) -> str:
    """Send a single image to Ollama for OCR."""
    resp = await _http_client.post(
        f"{OLLAMA_BASE}/api/chat",
        json={
            "model": OLLAMA_MODEL,
            "messages": [
                {
                    "role": "user",
                    "content": OCR_PROMPT,
                    "images": [image_b64],
                }
            ],
            "stream": False,
        },
    )
    resp.raise_for_status()
    result = resp.json()
    return result.get("message", {}).get("content", "")


def _postprocess(text: str) -> str:
    """Strip markdown fences and convert LaTeX to Unicode."""
    text = re.sub(r'^```\w*\n?', '', text.strip())
    text = re.sub(r'\n?```$', '', text.strip())
    text = _latex_to_unicode(text)
    return text.strip()


def _latex_to_unicode(text: str) -> str:
    """Replace LaTeX notation with Unicode characters using latex_unicode.json"""

    # 1. \textcircled{N} → ①②③...
    text = re.sub(
        r'\$\\textcircled\{(\d+)\}\$',
        lambda m: _CIRCLED.get(m.group(1), m.group(0)),
        text,
    )

    # 2. \frac{a}{b} → Unicode fraction (½ etc.) or a/b
    def _replace_frac(m):
        num, den = m.group(1), m.group(2)
        key = f"{num}/{den}"
        return _LATEX_FRACTIONS.get(key, f"{num}/{den}")

    text = re.sub(r'\$\\frac\{([^}]+)\}\{([^}]+)\}\$', _replace_frac, text)

    # 3. Simple $\command$ → Unicode (longest match first)
    for latex_cmd, unicode_char in _LATEX_SIMPLE:
        token = f"${latex_cmd}$"
        if token in text:
            text = text.replace(token, unicode_char)

    # 4. Remaining bare $\command$ patterns not in map — unwrap the $ delimiters
    text = re.sub(r'\$\\([a-zA-Z]+)\$', lambda m: '\\' + m.group(1), text)

    return text


def pdf_to_images(pdf_path: str, output_dir: Path) -> list[str]:
    """Convert PDF pages to images in the specified directory"""
    output_dir.mkdir(parents=True, exist_ok=True)
    filenames = []
    mat = fitz.Matrix(2.0, 2.0)
    doc = fitz.open(pdf_path)

    for page_num, page in enumerate(doc):
        pix = page.get_pixmap(matrix=mat)
        filename = f"page_{page_num + 1:03d}.png"
        pix.save(str(output_dir / filename))
        filenames.append(filename)

    doc.close()
    return filenames


def _safe_doc_path(doc_id: str, filename: str = "") -> Path:
    """Build a path inside UPLOAD_DIR/{doc_id} with traversal protection"""
    doc_dir = (UPLOAD_DIR / doc_id).resolve()
    if not str(doc_dir).startswith(str(UPLOAD_DIR.resolve())):
        raise HTTPException(403, "Invalid document ID")
    if filename:
        file_path = (doc_dir / filename).resolve()
        if not str(file_path).startswith(str(doc_dir)):
            raise HTTPException(403, "Invalid filename")
        return file_path
    return doc_dir


# --- Endpoints ---

@app.get("/", response_class=HTMLResponse)
async def root():
    """Serve frontend page"""
    with open("index.html", "r", encoding="utf-8") as f:
        return f.read()


@app.get("/api/status")
async def status():
    """Check service status"""
    ollama = await check_ollama()
    return {
        "status": "running",
        "model_loaded": ollama["model_loaded"],
        "device": "ollama",
        "gpu": {"name": f"Ollama ({OLLAMA_MODEL})"} if ollama["online"] else None,
    }


async def ensure_ollama_running() -> dict:
    """Start Ollama if not running, wait until ready"""
    ollama = await check_ollama()
    if ollama["online"]:
        return ollama

    logger.info("[ollama] Not running, attempting to start ollama serve...")
    try:
        subprocess.Popen(
            ["ollama", "serve"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            creationflags=subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0,
        )
        logger.info("[ollama] Popen launched, waiting for service...")
    except FileNotFoundError:
        logger.error("[ollama] 'ollama' command not found in PATH")
        raise HTTPException(500, "Ollama not found. Please install Ollama first.")
    except Exception as e:
        logger.error(f"[ollama] Failed to start: {e}", exc_info=True)
        raise HTTPException(500, f"Failed to start Ollama: {e}")

    for i in range(60):
        await asyncio.sleep(0.5)
        ollama = await check_ollama()
        if ollama["online"]:
            logger.info(f"[ollama] Started in {(i + 1) * 0.5:.1f}s")
            return ollama

    raise HTTPException(500, "Failed to start Ollama after 30s")


@app.post("/api/load-model")
async def load_model_endpoint():
    """Start Ollama if needed, then pre-warm model into GPU"""
    ollama = await ensure_ollama_running()

    if not ollama["model_loaded"]:
        raise HTTPException(500, f"Model '{OLLAMA_MODEL}' not found. Run: ollama pull {OLLAMA_MODEL}")

    try:
        t0 = time.time()
        await _http_client.post(
            f"{OLLAMA_BASE}/api/chat",
            json={"model": OLLAMA_MODEL, "messages": [{"role": "user", "content": "hi"}], "stream": False},
        )
        logger.info(f"[load_model] Warmup done: {time.time() - t0:.2f}s")
    except Exception as e:
        logger.warning(f"[load_model] Warmup failed: {e}")
    return {"success": True, "message": f"Model {OLLAMA_MODEL} loaded into GPU"}


ALLOWED_SUFFIXES = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.pdf'}


@app.post("/api/upload")
async def upload_files(files: list[UploadFile] = File(...)):
    """Upload one or more files. Multiple images become pages of one document.
    Always streams pages via SSE so the frontend gets incremental updates."""
    if not files:
        raise HTTPException(400, "No files provided")

    for f in files:
        if Path(f.filename).suffix.lower() not in ALLOWED_SUFFIXES:
            raise HTTPException(400, f"Unsupported file type: {f.filename}")

    # Read all file contents upfront (before entering the generator)
    file_data: list[tuple[str, str, bytes]] = []  # (filename, suffix, content)
    for f in files:
        suffix = Path(f.filename).suffix.lower()
        content = await f.read()
        file_data.append((f.filename, suffix, content))

    doc_id = str(uuid.uuid4())
    doc_dir = UPLOAD_DIR / doc_id
    doc_dir.mkdir(parents=True, exist_ok=True)

    # Display name
    if len(file_data) == 1:
        display_name = file_data[0][0]
    else:
        display_name = f"{len(file_data)} files"

    documents[doc_id] = {
        "filename": display_name,
        "pages": [],
        "created_at": datetime.now().isoformat(),
    }

    async def generate():
        page_num = 0

        yield f"data: {json.dumps({'type': 'init', 'doc_id': doc_id, 'filename': display_name})}\n\n"

        for fname, suffix, content in file_data:
            if suffix == ".pdf":
                # Save PDF, extract pages
                pdf_path = doc_dir / f"src_{uuid.uuid4().hex[:8]}.pdf"
                with open(pdf_path, "wb") as fp:
                    fp.write(content)

                mat = fitz.Matrix(2.0, 2.0)
                doc = fitz.open(str(pdf_path))
                for fitz_page in doc:
                    page_num += 1
                    pix = fitz_page.get_pixmap(matrix=mat)
                    img_name = f"page_{page_num:03d}.png"
                    pix.save(str(doc_dir / img_name))

                    page_info = {
                        "num": page_num,
                        "filename": img_name,
                        "image_url": f"/api/images/{doc_id}/{img_name}",
                        "ocr_text": None,
                        "ocr_regions": None,
                        "ocr_time": None,
                    }
                    documents[doc_id]["pages"].append(page_info)
                    yield f"data: {json.dumps({'type': 'page', 'page': page_info})}\n\n"
                    await asyncio.sleep(0)

                doc.close()
                pdf_path.unlink(missing_ok=True)
            else:
                # Single image
                page_num += 1
                img_name = f"page_{page_num:03d}{suffix}"
                with open(doc_dir / img_name, "wb") as fp:
                    fp.write(content)

                page_info = {
                    "num": page_num,
                    "filename": img_name,
                    "image_url": f"/api/images/{doc_id}/{img_name}",
                    "ocr_text": None,
                    "ocr_time": None,
                }
                documents[doc_id]["pages"].append(page_info)
                yield f"data: {json.dumps({'type': 'page', 'page': page_info})}\n\n"
                await asyncio.sleep(0)

        yield f"data: {json.dumps({'type': 'done', 'page_count': page_num})}\n\n"
        logger.info(f"[upload] {display_name} -> {doc_id}, {page_num} page(s)")

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.get("/api/images/{doc_id}/{filename}")
async def get_image(doc_id: str, filename: str):
    """Serve an uploaded page image"""
    file_path = _safe_doc_path(doc_id, filename)
    if not file_path.exists():
        raise HTTPException(404, "Image not found")
    return FileResponse(file_path, media_type="image/png")


@app.post("/api/ocr/{doc_id}/{page_num}")
async def ocr_single_page(doc_id: str, page_num: int, layout: bool = Query(True)):
    """Run OCR on a single page. Pass ?layout=false to skip layout detection."""
    if doc_id not in documents:
        raise HTTPException(404, "Document not found")

    doc = documents[doc_id]
    page = None
    for p in doc["pages"]:
        if p["num"] == page_num:
            page = p
            break
    if page is None:
        raise HTTPException(404, f"Page {page_num} not found")

    # If already OCR'd, return cached result
    if page["ocr_text"] is not None:
        return {
            "doc_id": doc_id,
            "page_num": page_num,
            "text": page["ocr_text"],
            "regions": page.get("ocr_regions") or [],
            "time": page["ocr_time"],
            "cached": True,
        }

    image_path = _safe_doc_path(doc_id, page["filename"])
    if not image_path.exists():
        raise HTTPException(404, "Image file not found")

    try:
        t0 = time.time()
        if layout:
            text, regions = await ocr_image_with_layout(str(image_path))
        else:
            img = Image.open(str(image_path)).convert("RGB")
            text = await _ocr_whole_image(img)
            img.close()
            regions = []
        elapsed = round(time.time() - t0, 2)

        page["ocr_text"] = text
        page["ocr_regions"] = regions
        page["ocr_time"] = elapsed

        return {
            "doc_id": doc_id,
            "page_num": page_num,
            "text": text,
            "regions": regions,
            "time": elapsed,
            "cached": False,
        }
    except httpx.HTTPStatusError as e:
        detail = e.response.text
        logger.error(f"[OCR] Ollama error: {detail}", exc_info=True)
        raise HTTPException(500, f"OCR failed: {detail}")
    except Exception as e:
        logger.error(f"[OCR] Error: {e}", exc_info=True)
        raise HTTPException(500, f"OCR failed: {e}")


@app.post("/api/ocr/{doc_id}/all")
async def ocr_all_pages(doc_id: str, layout: bool = Query(True)):
    """Run OCR on all pages of a document. Pass ?layout=false to skip layout detection."""
    if doc_id not in documents:
        raise HTTPException(404, "Document not found")

    doc = documents[doc_id]
    results = []

    for page in doc["pages"]:
        if page["ocr_text"] is not None:
            results.append({
                "page_num": page["num"],
                "text": page["ocr_text"],
                "regions": page.get("ocr_regions") or [],
                "time": page["ocr_time"],
                "cached": True,
            })
            continue

        image_path = _safe_doc_path(doc_id, page["filename"])
        try:
            t0 = time.time()
            if layout:
                text, regions = await ocr_image_with_layout(str(image_path))
            else:
                img = Image.open(str(image_path)).convert("RGB")
                text = await _ocr_whole_image(img)
                img.close()
                regions = []
            elapsed = round(time.time() - t0, 2)

            page["ocr_text"] = text
            page["ocr_regions"] = regions
            page["ocr_time"] = elapsed

            results.append({
                "page_num": page["num"],
                "text": text,
                "regions": regions,
                "time": elapsed,
                "cached": False,
            })
        except Exception as e:
            logger.error(f"[OCR] Page {page['num']} error: {e}", exc_info=True)
            results.append({
                "page_num": page["num"],
                "text": None,
                "regions": [],
                "time": None,
                "error": str(e),
            })

    return {
        "doc_id": doc_id,
        "filename": doc["filename"],
        "results": results,
    }


@app.delete("/api/documents/{doc_id}")
async def delete_document(doc_id: str):
    """Delete a document and its images"""
    if doc_id not in documents:
        raise HTTPException(404, "Document not found")

    doc_dir = _safe_doc_path(doc_id)
    if doc_dir.exists():
        shutil.rmtree(doc_dir, ignore_errors=True)

    del documents[doc_id]
    logger.info(f"[delete] Document {doc_id} removed")
    return {"success": True}


# --- DOCX Export ---

class _ExportPage(BaseModel):
    num: int
    text: str

class _ExportRequest(BaseModel):
    pages: list[_ExportPage]
    title: str | None = None


class _TableParser(HTMLParser):
    """Extract rows/cells from an HTML <table>."""
    def __init__(self):
        super().__init__()
        self.rows: list[list[str]] = []
        self._in_cell = False
        self._cell_text = ""
        self._current_row: list[str] = []
        self._is_header_row = False
        self.header_row_indices: set[int] = set()

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag == "tr":
            self._current_row = []
            self._is_header_row = False
        elif tag in ("td", "th"):
            self._in_cell = True
            self._cell_text = ""
            if tag == "th":
                self._is_header_row = True

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in ("td", "th"):
            self._in_cell = False
            self._current_row.append(self._cell_text.strip())
        elif tag == "tr":
            if self._current_row:
                if self._is_header_row:
                    self.header_row_indices.add(len(self.rows))
                self.rows.append(self._current_row)

    def handle_data(self, data):
        if self._in_cell:
            self._cell_text += data


def _parse_md_table(lines: list[str]) -> list[list[str]]:
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        stripped = line.strip()
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        # Skip separator row (---, :--:, etc.)
        if all(re.match(r'^[-:]+$', c) for c in cells if c):
            continue
        rows.append(cells)
    return rows


def _parse_ocr_text(text: str) -> list[dict]:
    """Parse OCR text into structured elements: headings, paragraphs, tables.
    Returns list of {type: 'heading'|'paragraph'|'table', ...}
    """
    elements = []

    # Split by HTML tables first
    parts = re.split(r'(<table[\s\S]*?</table>)', text, flags=re.IGNORECASE)

    for part in parts:
        if re.match(r'^<table[\s\S]*</table>$', part, re.IGNORECASE):
            # HTML table
            parser = _TableParser()
            parser.feed(part)
            if parser.rows:
                elements.append({
                    "type": "table",
                    "rows": parser.rows,
                    "header_rows": parser.header_row_indices,
                })
            continue

        # Process non-table text line by line
        lines = part.split('\n')
        md_table_buf = []

        def flush_md_table():
            if not md_table_buf:
                return
            rows = _parse_md_table(md_table_buf)
            if rows:
                elements.append({
                    "type": "table",
                    "rows": rows,
                    "header_rows": {0},  # first row is header in md tables
                })
            md_table_buf.clear()

        for line in lines:
            trimmed = line.strip()

            # Detect markdown table rows
            if '|' in trimmed and (trimmed.startswith('|') or re.search(r'\w\s*\|', trimmed)):
                md_table_buf.append(trimmed)
                continue

            if md_table_buf:
                flush_md_table()

            # Headings
            if trimmed.startswith('### '):
                elements.append({"type": "heading", "level": 3, "text": trimmed[4:]})
            elif trimmed.startswith('## '):
                elements.append({"type": "heading", "level": 2, "text": trimmed[3:]})
            elif trimmed.startswith('# '):
                elements.append({"type": "heading", "level": 1, "text": trimmed[2:]})
            elif trimmed:
                elements.append({"type": "paragraph", "text": trimmed})
            # Skip blank lines (they're just spacing)

        flush_md_table()

    return elements


def _build_docx(title: str, pages: list[_ExportPage]) -> io.BytesIO:
    """Build a real DOCX file from parsed OCR text."""
    doc = DocxDocument()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    # Set East Asian font (微软雅黑) for Chinese text
    rpr = style.element.get_or_add_rPr()
    ea_font = rpr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): '微软雅黑'})
    rpr.append(ea_font)

    multi_page = len(pages) > 1

    if multi_page:
        doc.add_heading(title, level=0)

    for idx, page in enumerate(pages):
        # Section break for each page after the first (each OCR page = one Word section)
        if multi_page and idx > 0:
            doc.add_section(WD_SECTION_START.NEW_PAGE)

        elements = _parse_ocr_text(page.text or '')

        for elem in elements:
            if elem["type"] == "heading":
                doc.add_heading(elem["text"], level=elem["level"])

            elif elem["type"] == "paragraph":
                para = doc.add_paragraph()
                text = elem["text"]
                # Simple bold/italic parsing
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
                for p in parts:
                    if p.startswith('**') and p.endswith('**'):
                        run = para.add_run(p[2:-2])
                        run.bold = True
                    elif p.startswith('*') and p.endswith('*') and len(p) > 2:
                        run = para.add_run(p[1:-1])
                        run.italic = True
                    else:
                        para.add_run(p)

            elif elem["type"] == "table":
                rows = elem["rows"]
                if not rows:
                    continue
                n_cols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=n_cols)
                tbl.style = 'Table Grid'

                header_rows = elem.get("header_rows", set())

                for i, row_data in enumerate(rows):
                    row = tbl.rows[i]
                    for j, cell_text in enumerate(row_data):
                        if j < n_cols:
                            cell = row.cells[j]
                            cell.text = cell_text
                            # Bold header cells
                            if i in header_rows:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.bold = True

    # Add page number footers (one per section = one per OCR page)
    if multi_page:
        for idx, section in enumerate(doc.sections):
            footer = section.footer
            footer.is_linked_to_previous = False
            para = footer.paragraphs[0]
            para.text = f"— {pages[idx].num} —"
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.post("/api/export/{doc_id}")
async def export_docx(doc_id: str, req: _ExportRequest):
    """Export document as a real DOCX file."""
    if doc_id not in documents:
        raise HTTPException(404, "Document not found")

    doc_meta = documents[doc_id]
    fallback = doc_meta.get("filename", "Document").replace(".pdf", "")
    title = req.title or fallback

    buf = _build_docx(title, req.pages)

    # RFC 5987 encoding for non-ASCII filenames
    safe_name = f"{title}.docx"
    encoded_name = quote(safe_name)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"
        },
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=3000)
